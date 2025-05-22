"""
Export utilities for Agentic Researcher

This module provides utilities for exporting research results to various formats
including PDF, DOCX, and markdown files.
"""

import os
import io
import logging
import base64
from typing import Dict, List, Any, Optional, Union
from datetime import datetime

# Import for PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, ListFlowable, ListItem
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available; PDF export will be limited")

# Import for DOCX generation
try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available; DOCX export will be limited")

# Import for Markdown to HTML conversion (for PDF generation)
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logging.warning("markdown not available; markdown conversion will be limited")

# Configure logging
logger = logging.getLogger(__name__)


class ExportManager:
    """
    Handles exporting research results to various formats

    Supported formats:
    - PDF
    - DOCX
    - Markdown
    - HTML
    """

    def __init__(self):
        """Initialize the export manager with default settings"""
        self.check_dependencies()

    def check_dependencies(self) -> Dict[str, bool]:
        """
        Check which export dependencies are available
        
        Returns:
            Dict of format names and availability
        """
        available = {
            "pdf": REPORTLAB_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "markdown": True,  # Always available
            "html": MARKDOWN_AVAILABLE,
        }
        
        logger.info(f"Export formats available: {', '.join([k for k, v in available.items() if v])}")
        return available

    def export_to_markdown(self, 
                          content: str,
                          filename: str = "research_results.md",
                          output_dir: str = "./exports") -> Dict[str, Any]:
        """
        Export content to a markdown file
        
        Args:
            content: Markdown content
            filename: Output filename
            output_dir: Directory for output
            
        Returns:
            Dict with success status and file path
        """
        try:
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Create full file path
            file_path = os.path.join(output_dir, filename)
            
            # Write markdown content to file
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Markdown export successful: {file_path}")
            return {
                "success": True,
                "format": "markdown",
                "file_path": file_path,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error exporting to markdown: {str(e)}")
            return {
                "success": False,
                "format": "markdown",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
    
    def export_to_html(self, 
                      content: str,
                      filename: str = "research_results.html",
                      output_dir: str = "./exports",
                      title: str = "Research Results") -> Dict[str, Any]:
        """
        Export markdown content to HTML
        
        Args:
            content: Markdown content
            filename: Output filename
            output_dir: Directory for output
            title: HTML document title
            
        Returns:
            Dict with success status and file path
        """
        if not MARKDOWN_AVAILABLE:
            return {
                "success": False,
                "format": "html",
                "error": "markdown library not available",
                "timestamp": datetime.now().isoformat()
            }
            
        try:
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Create full file path
            file_path = os.path.join(output_dir, filename)
            
            # Convert markdown to HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # Create a full HTML document
            html_document = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 24px;
            margin-bottom: 16px;
        }}
        h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: .3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: .3em; }}
        pre {{
            background-color: #f6f8fa;
            border-radius: 3px;
            padding: 16px;
            overflow: auto;
        }}
        code {{
            background-color: rgba(27, 31, 35, .05);
            border-radius: 3px;
            padding: 0.2em 0.4em;
            font-family: SFMono-Regular, Consolas, Liberation Mono, Menlo, monospace;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 16px;
        }}
        table, th, td {{
            border: 1px solid #dfe2e5;
        }}
        th, td {{
            padding: 12px 16px;
        }}
        th {{
            background-color: #f6f8fa;
        }}
        blockquote {{
            border-left: 4px solid #dfe2e5;
            padding: 0 16px;
            color: #6a737d;
            margin-left: 0;
            margin-right: 0;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="content">
        {html_content}
    </div>
    <footer>
        <p><small>Generated by Agentic Researcher on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</small></p>
    </footer>
</body>
</html>
"""
            
            # Write HTML content to file
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_document)
            
            logger.info(f"HTML export successful: {file_path}")
            return {
                "success": True,
                "format": "html",
                "file_path": file_path,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error exporting to HTML: {str(e)}")
            return {
                "success": False,
                "format": "html",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
    
    def export_to_pdf(self,
                     content: str,
                     filename: str = "research_results.pdf",
                     output_dir: str = "./exports",
                     title: str = "Research Results",
                     author: str = "Agentic Researcher") -> Dict[str, Any]:
        """
        Export markdown content to PDF using ReportLab
        
        Args:
            content: Markdown content
            filename: Output filename
            output_dir: Directory for output
            title: Document title
            author: Document author
            
        Returns:
            Dict with success status and file path
        """
        if not REPORTLAB_AVAILABLE or not MARKDOWN_AVAILABLE:
            return {
                "success": False,
                "format": "pdf",
                "error": "Required libraries not available: ReportLab and/or markdown",
                "timestamp": datetime.now().isoformat()
            }
            
        try:
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Create full file path
            file_path = os.path.join(output_dir, filename)
            
            # Convert markdown to HTML for better formatting
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # Setup document with metadata
            doc = SimpleDocTemplate(
                file_path,
                pagesize=A4,
                title=title,
                author=author
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='Code',
                fontName='Courier',
                fontSize=9,
                backColor=colors.lightgrey,
                spaceBefore=6,
                spaceAfter=6
            ))
            
            # Create document elements
            elements = []
            
            # Add title
            elements.append(Paragraph(title, styles['Title']))
            elements.append(Spacer(1, 12))
            
            # Add timestamp
            elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Italic']))
            elements.append(Spacer(1, 24))
            
            # Process HTML content into reportlab elements
            # This is a simplified approach, a more sophisticated parser would handle all HTML elements
            # Split content by major HTML tags and convert to appropriate elements
            current_text = ""
            in_code_block = False
            in_list = False
            list_items = []
            
            for line in html_content.split('\n'):
                if line.startswith('<h1>'):
                    if current_text:
                        elements.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    header_text = line.replace('<h1>', '').replace('</h1>', '')
                    elements.append(Paragraph(header_text, styles['Heading1']))
                    
                elif line.startswith('<h2>'):
                    if current_text:
                        elements.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    header_text = line.replace('<h2>', '').replace('</h2>', '')
                    elements.append(Paragraph(header_text, styles['Heading2']))
                    
                elif line.startswith('<h3>'):
                    if current_text:
                        elements.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    header_text = line.replace('<h3>', '').replace('</h3>', '')
                    elements.append(Paragraph(header_text, styles['Heading3']))
                    
                elif '<pre><code>' in line or line.startswith('<pre><code>'):
                    if current_text:
                        elements.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    in_code_block = True
                    code_text = line.replace('<pre><code>', '')
                    if '</code></pre>' in line:
                        code_text = code_text.replace('</code></pre>', '')
                        elements.append(Paragraph(code_text, styles['Code']))
                        in_code_block = False
                    else:
                        current_text = code_text
                        
                elif in_code_block and '</code></pre>' in line:
                    code_text = current_text + '\n' + line.replace('</code></pre>', '')
                    elements.append(Paragraph(code_text, styles['Code']))
                    current_text = ""
                    in_code_block = False
                    
                elif in_code_block:
                    current_text += '\n' + line
                    
                elif line.startswith('<ul>'):
                    if current_text:
                        elements.append(Paragraph(current_text, styles['Normal']))
                        current_text = ""
                    in_list = True
                    list_items = []
                    
                elif line.startswith('<li>') and in_list:
                    item_text = line.replace('<li>', '').replace('</li>', '')
                    list_items.append(ListItem(Paragraph(item_text, styles['Normal'])))
                    
                elif line.startswith('</ul>') and in_list:
                    elements.append(ListFlowable(list_items, bulletType='bullet'))
                    in_list = False
                    list_items = []
                    
                else:
                    if in_code_block:
                        current_text += '\n' + line
                    elif not in_list and line.strip() and not line.startswith('<'):
                        if current_text:
                            current_text += ' ' + line
                        else:
                            current_text = line
            
            # Add any remaining text
            if current_text:
                if in_code_block:
                    elements.append(Paragraph(current_text, styles['Code']))
                else:
                    elements.append(Paragraph(current_text, styles['Normal']))
            
            # Build the PDF document
            doc.build(elements)
            
            logger.info(f"PDF export successful: {file_path}")
            return {
                "success": True,
                "format": "pdf",
                "file_path": file_path,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {str(e)}")
            return {
                "success": False,
                "format": "pdf",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
    
    def export_to_docx(self,
                      content: str,
                      filename: str = "research_results.docx",
                      output_dir: str = "./exports",
                      title: str = "Research Results",
                      author: str = "Agentic Researcher") -> Dict[str, Any]:
        """
        Export markdown content to DOCX
        
        Args:
            content: Markdown content
            filename: Output filename
            output_dir: Directory for output
            title: Document title
            author: Document author
            
        Returns:
            Dict with success status and file path
        """
        if not DOCX_AVAILABLE or not MARKDOWN_AVAILABLE:
            return {
                "success": False,
                "format": "docx",
                "error": "Required libraries not available: python-docx and/or markdown",
                "timestamp": datetime.now().isoformat()
            }
            
        try:
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # Create full file path
            file_path = os.path.join(output_dir, filename)
            
            # Convert markdown to HTML for better parsing
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # Create a new document
            doc = docx.Document()
            
            # Set document properties
            doc.core_properties.title = title
            doc.core_properties.author = author
            
            # Add title
            doc.add_heading(title, level=0)
            
            # Add timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_run.italic = True
            
            # Add a line break
            doc.add_paragraph()
            
            # Process HTML content into docx elements
            # This is a simplified approach, a more sophisticated parser would handle all HTML elements
            current_text = ""
            in_code_block = False
            
            for line in html_content.split('\n'):
                if line.startswith('<h1>'):
                    if current_text:
                        doc.add_paragraph(current_text)
                        current_text = ""
                    header_text = line.replace('<h1>', '').replace('</h1>', '')
                    doc.add_heading(header_text, level=1)
                    
                elif line.startswith('<h2>'):
                    if current_text:
                        doc.add_paragraph(current_text)
                        current_text = ""
                    header_text = line.replace('<h2>', '').replace('</h2>', '')
                    doc.add_heading(header_text, level=2)
                    
                elif line.startswith('<h3>'):
                    if current_text:
                        doc.add_paragraph(current_text)
                        current_text = ""
                    header_text = line.replace('<h3>', '').replace('</h3>', '')
                    doc.add_heading(header_text, level=3)
                    
                elif '<pre><code>' in line or line.startswith('<pre><code>'):
                    if current_text:
                        doc.add_paragraph(current_text)
                        current_text = ""
                    in_code_block = True
                    code_text = line.replace('<pre><code>', '')
                    if '</code></pre>' in line:
                        code_text = code_text.replace('</code></pre>', '')
                        code_para = doc.add_paragraph(style='Normal')
                        code_run = code_para.add_run(code_text)
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(9)
                        in_code_block = False
                    else:
                        current_text = code_text
                        
                elif in_code_block and '</code></pre>' in line:
                    code_text = current_text + '\n' + line.replace('</code></pre>', '')
                    code_para = doc.add_paragraph(style='Normal')
                    code_run = code_para.add_run(code_text)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    current_text = ""
                    in_code_block = False
                    
                elif in_code_block:
                    current_text += '\n' + line
                    
                elif line.startswith('<ul>'):
                    if current_text:
                        doc.add_paragraph(current_text)
                        current_text = ""
                    
                elif line.startswith('<li>'):
                    item_text = line.replace('<li>', '').replace('</li>', '')
                    doc.add_paragraph(item_text, style='List Bullet')
                    
                elif line.startswith('</ul>'):
                    # Do nothing, list is already processed
                    pass
                    
                else:
                    if in_code_block:
                        current_text += '\n' + line
                    elif line.strip() and not line.startswith('<'):
                        if current_text:
                            current_text += ' ' + line
                        else:
                            current_text = line
            
            # Add any remaining text
            if current_text:
                if in_code_block:
                    code_para = doc.add_paragraph(style='Normal')
                    code_run = code_para.add_run(current_text)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                else:
                    doc.add_paragraph(current_text)
            
            # Save the document
            doc.save(file_path)
            
            logger.info(f"DOCX export successful: {file_path}")
            return {
                "success": True,
                "format": "docx",
                "file_path": file_path,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {str(e)}")
            return {
                "success": False,
                "format": "docx",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            }
    
    def export(self,
              content: str,
              format: str = "markdown",
              filename: Optional[str] = None,
              output_dir: str = "./exports",
              title: str = "Research Results",
              author: str = "Agentic Researcher") -> Dict[str, Any]:
        """
        Export content to the specified format
        
        Args:
            content: Content to export (typically markdown)
            format: Export format (pdf, docx, html, markdown)
            filename: Output filename (if None, generated from title and format)
            output_dir: Directory for output
            title: Document title
            author: Document author
            
        Returns:
            Dict with success status and file path
        """
        # Validate format
        format = format.lower()
        if format not in ["pdf", "docx", "html", "markdown", "md"]:
            return {
                "success": False,
                "error": f"Unsupported export format: {format}",
                "timestamp": datetime.now().isoformat()
            }
        
        # Normalize format (md -> markdown)
        if format == "md":
            format = "markdown"
        
        # Generate filename if not provided
        if not filename:
            safe_title = title.replace(" ", "_").lower()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}.{format}"
            if format == "markdown":
                filename = f"{safe_title}_{timestamp}.md"
        
        # Call appropriate export method
        if format == "pdf":
            return self.export_to_pdf(content, filename, output_dir, title, author)
        elif format == "docx":
            return self.export_to_docx(content, filename, output_dir, title, author)
        elif format == "html":
            return self.export_to_html(content, filename, output_dir, title)
        elif format == "markdown":
            return self.export_to_markdown(content, filename, output_dir)
        
        # Shouldn't reach here, but just in case
        return {
            "success": False,
            "error": "Unknown export format",
            "timestamp": datetime.now().isoformat()
        }
